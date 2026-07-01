"""全项目导出：分镜表链路 → Word 文档 + 素材 zip。"""

from __future__ import annotations

import asyncio
import json
import logging
import re
import shutil
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from urllib.parse import parse_qs, urlparse

from docx import Document
from docx.shared import Inches
from sqlalchemy.orm import Session

from db.session import SessionLocal
from models.export_job import ExportJob
from models.canvas_project import CanvasProject

logger = logging.getLogger(__name__)

_BACKEND_DIR = Path(__file__).resolve().parent.parent
UPLOADS_ROOT = _BACKEND_DIR / "uploads"
EXPORTS_DIR = UPLOADS_ROOT / "exports"

_SAFE_NAME = re.compile(r"[^\w\u4e00-\u9fff\-]+", re.UNICODE)


def _utcnow() -> datetime:
    return datetime.now(timezone.utc)


def _safe_zip_name(name: str, fallback: str = "export") -> str:
    text = (name or "").strip() or fallback
    cleaned = _SAFE_NAME.sub("_", text).strip("_")
    return cleaned[:80] or fallback


def _strip_media_ticket(url: str | None) -> str:
    raw = (url or "").strip()
    if not raw:
        return ""
    return raw.split("?", 1)[0]


def _resolve_upload_path(url: str | None) -> Path | None:
    """将 /api/uploads/... 或 /uploads/... 解析为本地文件路径。"""
    clean = _strip_media_ticket(url)
    if not clean:
        return None

    rel = clean
    if rel.startswith("/api/uploads/"):
        rel = rel[len("/api/uploads/") :]
    elif rel.startswith("/uploads/"):
        rel = rel[len("/uploads/") :]
    elif rel.startswith("http://") or rel.startswith("https://"):
        parsed = urlparse(rel)
        path = (parsed.path or "").split("?", 1)[0]
        if path.startswith("/api/uploads/"):
            rel = path[len("/api/uploads/") :]
        elif path.startswith("/uploads/"):
            rel = path[len("/uploads/") :]
        else:
            return None
    else:
        return None

    if not rel.startswith(("images/", "videos/")):
        return None

    candidates = [
        UPLOADS_ROOT / rel,
        _BACKEND_DIR / "uploads" / rel,
        Path(rel),
    ]
    for path in candidates:
        if path.is_file():
            return path.resolve()
    return None


def _resolve_view_path(url: str | None) -> Path | None:
    """尝试解析 /api/view?filename=... 为 Comfy 输出文件（若存在）。"""
    raw = (url or "").strip()
    if not raw or "filename=" not in raw:
        return None
    query = parse_qs(urlparse(raw).query)
    names = query.get("filename") or []
    if not names or not names[0]:
        return None
    filename = names[0].replace("\\", "/").split("/")[-1]
    subfolders = query.get("subfolder") or [""]
    subfolder = (subfolders[0] or "").strip().replace("\\", "/").strip("/")

    search_roots = [
        _BACKEND_DIR / "output",
        _BACKEND_DIR.parent / "output",
        UPLOADS_ROOT,
    ]
    for root in search_roots:
        if subfolder:
            candidate = root / subfolder / filename
            if candidate.is_file():
                return candidate.resolve()
        candidate = root / filename
        if candidate.is_file():
            return candidate.resolve()
    return None


def resolve_media_path(url: str | None) -> Path | None:
    path = _resolve_upload_path(url)
    if path:
        return path
    return _resolve_view_path(url)


def _load_canvas(project: CanvasProject) -> dict[str, Any]:
    try:
        data = json.loads(project.data or "{}")
    except json.JSONDecodeError as exc:
        raise ValueError("画布数据解析失败") from exc
    if not isinstance(data, dict):
        raise ValueError("画布数据格式无效")
    return data


def _node_by_id(nodes: list[dict], node_id: str) -> dict | None:
    for node in nodes:
        if node.get("id") == node_id:
            return node
    return None


def find_script_table_node(canvas: dict[str, Any], script_table_node_id: str) -> dict:
    nodes = canvas.get("nodes") or []
    node = _node_by_id(nodes, script_table_node_id)
    if not node or node.get("type") != "script-table":
        raise ValueError("分镜表节点不存在")
    return node


def find_outline_node(
    canvas: dict[str, Any],
    script_table_node: dict,
) -> dict | None:
    nodes = canvas.get("nodes") or []
    edges = canvas.get("edges") or []
    script_id = script_table_node.get("id")
    script_data = script_table_node.get("data") or {}

    outline_id = script_data.get("sourceOutlineId")
    if outline_id:
        node = _node_by_id(nodes, outline_id)
        if node and node.get("type") == "outline":
            return node

    for edge in edges:
        if edge.get("target") == script_id:
            src_id = edge.get("source")
            node = _node_by_id(nodes, src_id)
            if node and node.get("type") == "outline":
                return node

    for node in nodes:
        if node.get("type") != "outline":
            continue
        data = node.get("data") or {}
        if data.get("linkedScriptTableId") == script_id:
            return node
    return None


def _outline_export_paragraphs(outline_node: dict | None) -> list[str]:
    if not outline_node:
        return []
    data = outline_node.get("data") or {}
    lines: list[str] = []
    title = (data.get("title") or "").strip()
    if title:
        lines.append(title)
        lines.append("")

    scenes = data.get("scenes") or []
    if not scenes and data.get("versions"):
        versions = data.get("versions") or []
        idx = int(data.get("selectedVersionIndex") or 0)
        if versions and 0 <= idx < len(versions):
            scenes = versions[idx].get("scenes") or []
            if not title:
                vt = (versions[idx].get("title") or "").strip()
                if vt:
                    lines.insert(0, vt)
                    lines.insert(1, "")

    for scene in scenes:
        heading = (scene.get("title") or "").strip()
        if heading:
            lines.append(heading)
        for key, label in (
            ("location", "地点"),
            ("time", "时间"),
            ("characters", "人物"),
            ("mood", "氛围"),
        ):
            value = (scene.get(key) or "").strip()
            if value:
                lines.append(f"{label}：{value}")
        body = (scene.get("content") or "").strip()
        if body:
            lines.append(body)
        lines.append("")
    return lines


def _row_storyboard_url(row: dict, nodes: list[dict] | None = None) -> str | None:
    direct = row.get("directResultUrl") or row.get("resultUrl")
    if direct:
        return direct
    beat_id = row.get("beatCardNodeId")
    if beat_id and nodes:
        beat = _node_by_id(nodes, beat_id)
        if beat and beat.get("type") == "script-beat-card":
            for kf in (beat.get("data") or {}).get("keyframes") or []:
                url = kf.get("resultUrl")
                if url:
                    return url
    keyframes = row.get("keyframes") or []
    for kf in keyframes:
        url = kf.get("resultUrl")
        if url:
            return url
    return None


def _row_video_url(row: dict, nodes: list[dict]) -> str | None:
    for vid_id in (row.get("directVideoGenNodeId"), row.get("videoGenNodeId")):
        if not vid_id:
            continue
        node = _node_by_id(nodes, vid_id)
        if not node or node.get("type") != "video-gen":
            continue
        data = node.get("data") or {}
        if data.get("status") == "failed":
            continue
        if data.get("status") != "completed":
            continue
        url = data.get("videoUrl")
        if url:
            return url
    return None


def _row_prompt(row: dict) -> str:
    return (row.get("prompt") or row.get("description") or "").strip()


def _format_shot_index(shot_number: int | None, fallback: int) -> str:
    n = shot_number if shot_number is not None else fallback
    try:
        num = int(n)
    except (TypeError, ValueError):
        num = fallback
    return f"{num:02d}"


def _build_docx(
    *,
    project_name: str,
    outline_lines: list[str],
    shots: list[dict[str, Any]],
    temp_dir: Path,
) -> Path:
    doc = Document()
    doc.add_heading(project_name or "未命名项目", level=0)
    doc.add_paragraph(f"导出时间：{_utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_page_break()

    doc.add_heading("剧本 / 大纲", level=1)
    if outline_lines:
        for line in outline_lines:
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("（无关联大纲）")

    doc.add_page_break()
    doc.add_heading("分镜表", level=1)

    for idx, shot in enumerate(shots, start=1):
        shot_label = shot.get("label") or f"镜头 {shot.get('shot_number', idx)}"
        doc.add_heading(shot_label, level=2)
        duration = shot.get("duration")
        if duration is not None:
            doc.add_paragraph(f"时长：{duration}s")

        prompt = shot.get("prompt") or ""
        if prompt:
            doc.add_paragraph(prompt)

        image_note = shot.get("image_note")
        image_path: Path | None = shot.get("image_path")
        if image_path and image_path.is_file():
            try:
                doc.add_picture(str(image_path), width=Inches(5.5))
            except Exception:
                doc.add_paragraph("（分镜图嵌入失败）")
        elif image_note:
            doc.add_paragraph(image_note)

        video_note = shot.get("video_note")
        if video_note:
            doc.add_paragraph(video_note)

        doc.add_paragraph("")

    docx_path = temp_dir / f"{_safe_zip_name(project_name)}_export.docx"
    doc.save(docx_path)
    return docx_path


def _collect_shot_assets(
    rows: list[dict],
    nodes: list[dict],
    temp_dir: Path,
) -> list[dict[str, Any]]:
    images_dir = temp_dir / "images"
    videos_dir = temp_dir / "videos"
    images_dir.mkdir(parents=True, exist_ok=True)
    videos_dir.mkdir(parents=True, exist_ok=True)

    shots: list[dict[str, Any]] = []
    indexed_rows = list(enumerate(rows))
    sorted_rows = sorted(
        indexed_rows,
        key=lambda item: (
            int(item[1].get("shotNumber") or 0),
            item[0],
        ),
    )

    for idx, (_, row) in enumerate(sorted_rows, start=1):
        shot_num = row.get("shotNumber", idx)
        shot_index = _format_shot_index(shot_num, idx)
        prompt = _row_prompt(row)
        duration = row.get("duration")

        image_url = _row_storyboard_url(row, nodes)
        image_path = resolve_media_path(image_url)
        image_note = None
        image_copy: Path | None = None

        if image_path:
            ext = image_path.suffix or ".png"
            dest = images_dir / f"shot_{shot_index}{ext}"
            shutil.copy2(image_path, dest)
            image_copy = dest
        elif image_url:
            image_note = "（分镜图文件不可用）"
        else:
            image_note = "（图像生成中）"

        video_url = _row_video_url(row, nodes)
        video_path = resolve_media_path(video_url)
        video_note = None
        if video_path:
            ext = video_path.suffix or ".mp4"
            dest = videos_dir / f"shot_{shot_index}{ext}"
            shutil.copy2(video_path, dest)
        elif video_url:
            video_note = "（镜头视频文件不可用）"
        elif row.get("videoGenNodeId"):
            video_note = "（视频生成中）"

        shots.append(
            {
                "shot_number": shot_num,
                "label": f"镜头 {shot_num}（{duration or '?'}s）",
                "duration": duration,
                "prompt": prompt,
                "image_path": image_copy,
                "image_note": image_note,
                "video_note": video_note,
            }
        )
    return shots


def execute_export_job(export_job_id: str) -> None:
    db = SessionLocal()
    temp_dir: Path | None = None
    try:
        job = db.get(ExportJob, export_job_id)
        if not job:
            return

        job.status = "processing"
        db.commit()

        project = db.get(CanvasProject, job.project_id)
        if not project:
            raise ValueError("项目不存在")

        canvas = _load_canvas(project)
        script_node = find_script_table_node(canvas, job.script_table_node_id)
        outline_node = find_outline_node(canvas, script_node)
        rows = (script_node.get("data") or {}).get("rows") or []
        nodes = canvas.get("nodes") or []

        EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
        temp_dir = Path(tempfile.mkdtemp(prefix="export_job_"))

        shots = _collect_shot_assets(rows, nodes, temp_dir)
        outline_lines = _outline_export_paragraphs(outline_node)
        docx_path = _build_docx(
            project_name=project.name,
            outline_lines=outline_lines,
            shots=shots,
            temp_dir=temp_dir,
        )

        zip_name = f"{_safe_zip_name(project.name)}_{export_job_id[:8]}.zip"
        zip_path = EXPORTS_DIR / zip_name
        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.write(docx_path, arcname=docx_path.name)
            images_dir = temp_dir / "images"
            videos_dir = temp_dir / "videos"
            if images_dir.is_dir():
                for file in sorted(images_dir.iterdir()):
                    if file.is_file():
                        zf.write(file, arcname=f"images/{file.name}")
            if videos_dir.is_dir():
                for file in sorted(videos_dir.iterdir()):
                    if file.is_file():
                        zf.write(file, arcname=f"videos/{file.name}")

        rel_path = f"exports/{zip_name}"
        job.status = "completed"
        job.file_path = rel_path
        job.error_message = None
        db.commit()
        logger.info("export job completed: %s -> %s", export_job_id, rel_path)
    except Exception as exc:
        logger.exception("export job failed: %s", export_job_id)
        db.rollback()
        job = db.get(ExportJob, export_job_id)
        if job:
            job.status = "failed"
            job.error_message = str(exc)[:2000]
            db.commit()
    finally:
        if temp_dir and temp_dir.is_dir():
            shutil.rmtree(temp_dir, ignore_errors=True)
        db.close()


async def run_export_job(export_job_id: str) -> None:
    await asyncio.to_thread(execute_export_job, export_job_id)


def create_export_job_record(
    db: Session,
    *,
    project_id: str,
    script_table_node_id: str,
    user_id: int,
) -> ExportJob:
    from models.export_job import new_export_job_id

    job = ExportJob(
        id=new_export_job_id(),
        project_id=project_id,
        script_table_node_id=script_table_node_id,
        status="pending",
        created_by=user_id,
    )
    db.add(job)
    db.commit()
    db.refresh(job)
    return job


def get_export_job_for_user(db: Session, export_id: str, user_id: int) -> ExportJob:
    from fastapi import HTTPException

    from services.canvas_access import get_accessible_project

    from models import User

    job = db.get(ExportJob, export_id)
    if not job:
        raise HTTPException(status_code=404, detail="导出任务不存在")
    user = db.get(User, user_id)
    if not user:
        raise HTTPException(status_code=401, detail="未登录")
    get_accessible_project(db, user, job.project_id)
    return job


def export_job_to_dict(job: ExportJob) -> dict[str, Any]:
    download_url = None
    if job.status == "completed" and job.file_path:
        download_url = f"/api/exports/{job.id}/download"
    return {
        "id": job.id,
        "project_id": job.project_id,
        "script_table_node_id": job.script_table_node_id,
        "status": job.status,
        "file_path": job.file_path,
        "error_message": job.error_message,
        "download_url": download_url,
        "created_at": job.created_at.isoformat() if job.created_at else None,
    }
